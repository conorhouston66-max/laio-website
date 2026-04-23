from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.1)
section.right_margin = Inches(1.1)

# ── Colour palette
BLACK   = RGBColor(0x0F, 0x0F, 0x0F)
AMBER   = RGBColor(0xC8, 0x96, 0x0C)
MUTED   = RGBColor(0x55, 0x55, 0x55)
LIGHT   = RGBColor(0x99, 0x99, 0x99)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
CREAM   = RGBColor(0xF7, 0xF3, 0xEC)

def set_font(run, size=11, bold=False, color=BLACK, font='Calibri'):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=22, bold=True, color=BLACK, font='Calibri')
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=BLACK, font='Calibri')
    return p

def heading3(text, color=AMBER):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=color, font='Calibri')
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_font(run, size=10.5, color=MUTED)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' ')
        set_font(r1, size=10.5, bold=True, color=BLACK)
        r2 = p.add_run(text)
        set_font(r2, size=10.5, color=MUTED)
    else:
        run = p.add_run(text)
        set_font(run, size=10.5, color=MUTED)
    return p

def checkbox(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run('☐  ' + bold_prefix + '  ')
        set_font(r1, size=10.5, bold=True, color=BLACK)
        r2 = p.add_run(text)
        set_font(r2, size=10.5, color=MUTED)
    else:
        run = p.add_run('☐  ' + text)
        set_font(run, size=10.5, color=MUTED)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 80)
    set_font(run, size=8, color=LIGHT)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=WHITE)
        # Shade header cell
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F0F0F')
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            set_font(run, size=10, color=MUTED if ci > 0 else BLACK)
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F7F3EC')
                tcPr.append(shd)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('Lead Acquisition')
set_font(r, size=28, bold=True, color=BLACK, font='Calibri')

p2 = doc.add_paragraph()
r2 = p2.add_run('30-Day SEO & AEO Plan')
set_font(r2, size=18, bold=False, color=AMBER, font='Calibri')

p3 = doc.add_paragraph()
r3 = p3.add_run('April 2026 — May 2026')
set_font(r3, size=11, color=MUTED)

doc.add_paragraph()
body('Goal: Establish leadacquisition.io as the go-to result — in Google, Bing, and AI search engines — when a B2B decision-maker searches for a cold email agency, SDR outsourcing, or outbound lead generation. This plan covers technical SEO, content, citations, link building, and AEO (Answer Engine Optimisation for ChatGPT, Claude, Perplexity).')
divider()

# ══════════════════════════════════════════════════════════════
# WHERE WE START
# ══════════════════════════════════════════════════════════════
heading1('Where We Start')
body('Assets already in place at day 0:')
bullet('Website live at leadacquisition.io via Netlify with auto-deploy from GitHub')
bullet('JSON-LD schema markup: Organization, ProfessionalService, FAQPage, WebSite')
bullet('Hidden entity statement for AI crawler ingestion')
bullet('Meta title updated: "Lead Acquisition | B2B Cold Email Agency | Outbound Lead Generation"')
bullet('Daily automated blog posts running at 8am — 15 commercial-intent keywords queued')
bullet('Beehiiv RSS feed pulling live into homepage blog section')
bullet('First blog published: "Signal-Based Outbound: The B2B Guide to Reaching Buyers at the Right Moment"')

divider()

# ══════════════════════════════════════════════════════════════
# WEEK 1
# ══════════════════════════════════════════════════════════════
heading1('Week 1 — Technical Foundation & Indexing (Days 1–7)')
body('Nothing ranks if it is not indexed. Week 1 locks in the technical foundation and gets the site into every search index that matters including the ones that feed AI engines.')

heading2('Priority Actions')

heading3('1. Submit to Google Search Console')
checkbox('Go to search.google.com/search-console', bold_prefix='Action:')
checkbox('Add property → Domain → verify via DNS TXT record in Netlify DNS settings', bold_prefix='Action:')
checkbox('Submit sitemap: leadacquisition.io/sitemap.xml (create this — see below)', bold_prefix='Action:')
checkbox('Request indexing on homepage and all 3 existing blog posts', bold_prefix='Action:')

heading3('2. Create a Sitemap')
body('A sitemap tells Google every page that exists. Without it, blog posts may take weeks to be discovered.', indent=True)
checkbox('Create sitemap.xml in root folder (Claude can do this — just ask)', bold_prefix='Action:')
checkbox('Submit to Google Search Console', bold_prefix='Action:')
checkbox('Submit to Bing Webmaster Tools at bing.com/webmasters (critical — ChatGPT searches via Bing)', bold_prefix='Action:')

heading3('3. Submit to Bing Webmaster Tools')
body('ChatGPT uses Bing as its search index. If you are not in Bing, you are invisible to ChatGPT recommendations.', indent=True)
checkbox('Go to bing.com/webmasters', bold_prefix='Action:')
checkbox('Import your site from Google Search Console (one-click import)', bold_prefix='Action:')
checkbox('Submit sitemap', bold_prefix='Action:')

heading3('4. Set Up Crunchbase Company Profile')
body('Crunchbase is one of the most frequently cited sources when AI models recommend B2B vendors. Free to create.', indent=True)
checkbox('Go to crunchbase.com → Add your company', bold_prefix='Action:')
checkbox('Fill in: company name, description (use exact phrase "B2B cold email agency"), founded date, founder name, website, HQ location', bold_prefix='Action:')
checkbox('Add services: "Cold Email", "B2B Lead Generation", "Outbound Sales", "SDR Outsourcing"', bold_prefix='Action:')

heading3('5. Create a Clutch Profile')
body('Clutch is the most cited B2B services directory in AI search results. It also drives direct inbound.', indent=True)
checkbox('Go to clutch.co → Join as a provider', bold_prefix='Action:')
checkbox('Category: "Email Marketing", "Lead Generation", "Sales Outsourcing"', bold_prefix='Action:')
checkbox('Complete all fields. Ask 2-3 past clients to leave a review (even short ones — Clutch weight reviews heavily)', bold_prefix='Action:')

doc.add_paragraph()
heading2('Week 1 Deliverables Checklist')
add_table(
    ['Item', 'Owner', 'Status'],
    [
        ['Google Search Console set up + sitemap submitted', 'You', '☐'],
        ['Bing Webmaster Tools set up', 'You', '☐'],
        ['Crunchbase profile live', 'You', '☐'],
        ['Clutch profile live', 'You', '☐'],
        ['sitemap.xml created and live', 'Claude', '☐'],
        ['All blog posts indexed in Google', 'Auto', '☐'],
    ],
    col_widths=[3.2, 1.0, 0.9]
)

divider()

# ══════════════════════════════════════════════════════════════
# WEEK 2
# ══════════════════════════════════════════════════════════════
heading1('Week 2 — Content Velocity & On-Page SEO (Days 8–14)')
body('Week 2 is about getting content volume up fast. The daily blog routine is running, but you also need to optimise existing pages and build topical authority around your core keywords.')

heading2('Priority Actions')

heading3('1. Let the Daily Blog Run — Review Each Post')
body('The automated task publishes one post per day from the commercial-intent queue. Your job is to review each post in the morning and optionally copy it into Beehiiv to send as a newsletter to your subscriber list.', indent=True)
body('Posts going live this week:', indent=True)
bullet('"outsource cold email" — Day 8')
bullet('"cold email agency vs in-house" — Day 9')
bullet('"signs you need a cold email agency" — Day 10')
bullet('"SDR outsourcing" — Day 11')
bullet('"done for you cold email" — Day 12')
bullet('"b2b appointment setting" — Day 13')
bullet('"outbound sales outsourcing" — Day 14')

heading3('2. Add Internal Links Across Blog Posts')
body('Internal links pass authority between pages and help Google understand your site structure. Once you have 5+ posts live, ask Claude to audit internal linking and add cross-links between relevant posts.', indent=True)

heading3('3. Create a LinkedIn Company Page')
body('AI models check LinkedIn for entity verification. It also drives direct traffic and backlinks.', indent=True)
checkbox('Create company page: linkedin.com/company/lead-acquisition', bold_prefix='Action:')
checkbox('Add description using exact phrase "B2B cold email agency" and "outbound lead generation"', bold_prefix='Action:')
checkbox('Link back to leadacquisition.io', bold_prefix='Action:')
checkbox('Post each new blog article as a LinkedIn post (drives traffic + signals freshness)', bold_prefix='Action:')

heading3('4. Add Open Graph Meta Tags to All Pages')
body('OG tags control how your pages look when shared on LinkedIn, Twitter, Slack. A properly formatted preview drives click-through. Ask Claude to add these to all pages.', indent=True)
checkbox('Ask Claude: "Add Open Graph meta tags to index.html and all blog posts"', bold_prefix='Action:')

heading3('5. Create a Google Business Profile')
body('Even for a remote agency, a Google Business Profile helps with entity recognition and local pack inclusion. Use your registered business address.', indent=True)
checkbox('Go to business.google.com → Add your business', bold_prefix='Action:')
checkbox('Category: "Marketing Agency" or "Lead Generation Service"', bold_prefix='Action:')
checkbox('Add website, description, and photos', bold_prefix='Action:')

doc.add_paragraph()
heading2('Week 2 Deliverables Checklist')
add_table(
    ['Item', 'Owner', 'Status'],
    [
        ['7 new blog posts live (auto)', 'Claude (auto)', '☐'],
        ['LinkedIn company page created', 'You', '☐'],
        ['Open Graph tags added to all pages', 'Claude', '☐'],
        ['Google Business Profile live', 'You', '☐'],
        ['Each new post shared on LinkedIn', 'You', '☐'],
        ['Internal linking audit done', 'Claude', '☐'],
    ],
    col_widths=[3.2, 1.0, 0.9]
)

divider()

# ══════════════════════════════════════════════════════════════
# WEEK 3
# ══════════════════════════════════════════════════════════════
heading1('Week 3 — Link Building & Citations (Days 15–21)')
body('Links from other sites are still the primary ranking factor in 2026. Week 3 is about getting mentioned on authoritative pages that AI models already trust and cite.')

heading2('Priority Actions')

heading3('1. Guest Post Outreach — Target Existing "Best Cold Email Agency" Listicles')
body('This is the highest-leverage link building activity available. Find articles already ranking for "best cold email agencies 2026" and reach out to get Lead Acquisition added to the list. These pages already rank, already get cited by AI, and a mention there immediately puts you in front of buyers.', indent=True)
body('Target list (search Google for "best cold email agencies 2026" and identify the top 10 results):', indent=True)
bullet('Email the author/editor of each article')
bullet('Subject: "Addition for your cold email agency roundup"')
bullet('Body: Brief pitch on what makes Lead Acquisition different (parallel campaign testing, signal-based targeting, speed to first meetings). Offer a brief case study or stat.')
bullet('Aim for 3-5 inclusions in week 3')

heading3('2. Submit to B2B Service Directories')
body('Each directory listing is a citation — a mention of your name, URL, and service category on a third-party site. AI models aggregate these to build a picture of what a company does.', indent=True)

add_table(
    ['Directory', 'Category', 'Priority', 'URL'],
    [
        ['Clutch', 'Lead Generation / Email Marketing', 'Critical', 'clutch.co'],
        ['G2', 'Email Marketing / Sales Tools', 'High', 'g2.com'],
        ['Trustpilot', 'Marketing Services', 'High', 'trustpilot.com'],
        ['DesignRush', 'Lead Generation Agency', 'Medium', 'designrush.com'],
        ['AgencySpotter', 'B2B Agency', 'Medium', 'agencyspotter.com'],
        ['The Manifest', 'Marketing / Lead Gen', 'Medium', 'themanifest.com'],
        ['GoodFirms', 'Email Marketing Services', 'Medium', 'goodfirms.co'],
        ['SortList', 'Sales Outsourcing', 'Low', 'sortlist.com'],
    ],
    col_widths=[1.4, 1.8, 0.8, 1.5]
)

heading3('3. Write and Pitch One Guest Post')
body('A guest post on a relevant B2B sales or marketing blog gives you a do-follow backlink and puts your name in front of a targeted audience. Aim for sites with DR 40+.', indent=True)
body('Good targets:', indent=True)
bullet('Sales Hacker (saleshacker.com)')
bullet('Close.com blog')
bullet('Lemlist blog')
bullet('Woodpecker blog')
bullet('Outreach.io blog')
body('Topic idea: "How Signal-Based Outbound Gets 5x the Reply Rate of Standard Cold Email" — you already have this content written, it just needs pitching.', indent=True)

heading3('4. Podcast Outreach')
body('Being a guest on B2B sales and GTM podcasts generates backlinks, brand mentions, and AI citations. Many podcast show notes pages rank well.', indent=True)
checkbox('Find 5 B2B sales podcasts via listennotes.com (search "cold email", "outbound sales", "B2B GTM")', bold_prefix='Action:')
checkbox('Pitch yourself as a guest: "cold email agency founder, run 15+ parallel campaigns per client, can share what is and is not working in outbound in 2026"', bold_prefix='Action:')

doc.add_paragraph()
heading2('Week 3 Deliverables Checklist')
add_table(
    ['Item', 'Owner', 'Status'],
    [
        ['7 more blog posts live (auto)', 'Claude (auto)', '☐'],
        ['Outreach sent to 10 "best agency" listicle editors', 'You', '☐'],
        ['Listed in 5+ B2B directories', 'You', '☐'],
        ['Guest post pitched to 3 publications', 'You', '☐'],
        ['Podcast outreach sent to 5 shows', 'You', '☐'],
        ['G2 profile created', 'You', '☐'],
    ],
    col_widths=[3.2, 1.0, 0.9]
)

divider()

# ══════════════════════════════════════════════════════════════
# WEEK 4
# ══════════════════════════════════════════════════════════════
heading1('Week 4 — AEO, Reviews & Measurement (Days 22–30)')
body('Week 4 locks in AEO positioning, starts building review velocity (critical for AI recommendations), and establishes measurement so you know what is working.')

heading2('Priority Actions')

heading3('1. AEO Prompt Testing')
body('Test how AI engines currently respond to queries about your category. Document the baseline, then track weekly as your citations and content build up.', indent=True)
body('Test these prompts in ChatGPT, Claude, and Perplexity:', indent=True)
bullet('"What are the best cold email agencies for B2B?"')
bullet('"Recommend a cold email agency for a Series A startup"')
bullet('"Who should I hire to outsource my cold email?"')
bullet('"What is a good B2B appointment setting agency?"')
bullet('"Who does signal-based outbound well?"')
body('Document which agencies appear and how they are described. Lead Acquisition appearing is the goal. If not appearing yet, note which sites the AI cites — those are your link-building targets.', indent=True)

heading3('2. Get 3-5 Reviews on Clutch')
body('Clutch reviews are among the most cited social proof signals in AI recommendations for agencies. One review from a real client outweighs months of content.', indent=True)
checkbox('Email past clients with a direct Clutch review link', bold_prefix='Action:')
checkbox('Make it easy — offer to write a draft they can edit and submit', bold_prefix='Action:')
checkbox('Also request Google Business Profile reviews', bold_prefix='Action:')

heading3('3. Create a "Case Studies" Page')
body('AI models look for evidence that you get results. A dedicated case study page with specific numbers (X meetings in Y days, Z% reply rate) gives AI something concrete to cite when recommending you.', indent=True)
checkbox('Ask Claude to build a dedicated /case-studies page with 2-3 anonymised results', bold_prefix='Action:')

heading3('4. Build the Keyword Tracking Dashboard')
body('By day 22 you have 14+ pieces of content live. Time to start tracking which ones are ranking and driving traffic.', indent=True)
body('Free tools to set up:', indent=True)
bullet('Google Search Console — track impressions and clicks per page and keyword')
bullet('Bing Webmaster Tools — same for Bing')
bullet('Ahrefs Webmaster Tools (free) — ahrefs.com/webmaster-tools — keyword positions and backlink tracking')

heading3('5. Refresh and Expand Best-Performing Posts')
body('By day 28 you will have data from Search Console on which posts are getting impressions. Take the top 2-3 and ask Claude to expand them — add more sections, update statistics, improve internal linking. Google rewards freshness and depth.', indent=True)

doc.add_paragraph()
heading2('Week 4 Deliverables Checklist')
add_table(
    ['Item', 'Owner', 'Status'],
    [
        ['AEO prompt tests run and documented', 'You', '☐'],
        ['3+ Clutch reviews collected', 'You', '☐'],
        ['Case studies page live', 'Claude', '☐'],
        ['Google Search Console tracking set up', 'You', '☐'],
        ['Ahrefs Webmaster Tools set up', 'You', '☐'],
        ['Top 2 posts expanded and refreshed', 'Claude', '☐'],
        ['7 more blog posts live (auto)', 'Claude (auto)', '☐'],
    ],
    col_widths=[3.2, 1.0, 0.9]
)

divider()

# ══════════════════════════════════════════════════════════════
# FULL 30-DAY OVERVIEW
# ══════════════════════════════════════════════════════════════
heading1('Full 30-Day Overview')

add_table(
    ['Week', 'Focus', 'Key Actions', 'Owner'],
    [
        ['Week 1\nDays 1-7', 'Technical\nFoundation', 'Search Console, Bing WMT, Sitemap, Crunchbase, Clutch', 'You + Claude'],
        ['Week 2\nDays 8-14', 'Content &\nOn-Page', 'Daily blogs live, LinkedIn page, OG tags, Google Business, internal links', 'You + Claude'],
        ['Week 3\nDays 15-21', 'Link Building\n& Citations', 'Listicle outreach, B2B directories, guest post pitch, podcast outreach', 'You'],
        ['Week 4\nDays 22-30', 'AEO, Reviews\n& Measure', 'Prompt testing, Clutch reviews, case studies, Search Console tracking', 'You + Claude'],
    ],
    col_widths=[1.0, 1.0, 2.8, 1.0]
)

divider()

# ══════════════════════════════════════════════════════════════
# KEYWORD TARGETS
# ══════════════════════════════════════════════════════════════
heading1('Target Keywords — 30-Day Ranking Priorities')
body('These are the keywords where ranking in the top 10 means qualified inbound. Ordered by commercial intent and rankability for a new domain.')

add_table(
    ['Keyword', 'Intent', 'Difficulty', 'Target Page'],
    [
        ['outsource cold email', 'Commercial', 'Medium', 'Blog post (auto — Day 8)'],
        ['cold email agency vs in-house', 'Commercial', 'Low', 'Blog post (auto — Day 9)'],
        ['done for you cold email', 'Commercial', 'Low', 'Blog post (auto — Day 12)'],
        ['SDR outsourcing', 'Commercial', 'Low-Med', 'Blog post (auto — Day 11)'],
        ['b2b appointment setting', 'Commercial', 'Medium', 'Blog post (auto — Day 13)'],
        ['signs you need a cold email agency', 'Pain-aware', 'Very Low', 'Blog post (auto — Day 10)'],
        ['cold email agency results', 'Commercial', 'Low', 'Blog post (auto — Day 21)'],
        ['signal-based outbound', 'Informational', 'Very Low', 'LIVE — blog.leadacquisition.io'],
        ['outbound sales for series a startup', 'Commercial', 'Very Low', 'Blog post (auto — Day 25)'],
        ['how to choose a cold email agency', 'Commercial', 'Low', 'Blog post (auto — Day 18)'],
    ],
    col_widths=[2.2, 1.0, 0.9, 2.0]
)

divider()

# ══════════════════════════════════════════════════════════════
# AEO STRATEGY
# ══════════════════════════════════════════════════════════════
heading1('AEO Strategy — Getting Recommended by AI Engines')
body('When someone asks ChatGPT, Claude, or Perplexity "what is the best cold email agency?", here is what drives the response:')

add_table(
    ['Signal', 'What It Does', 'Status'],
    [
        ['JSON-LD Schema (Organization + Service + FAQ)', 'Tells AI crawlers exactly what you are and what you offer', 'LIVE'],
        ['Entity statement (hidden body text)', 'Plain-text crawlable description loaded with target phrases', 'LIVE'],
        ['Clutch profile with reviews', 'Most cited B2B agency directory in AI responses', 'TO DO'],
        ['Crunchbase company page', 'Frequently cited for entity verification', 'TO DO'],
        ['Bing indexing', 'ChatGPT searches via Bing — must be indexed', 'TO DO'],
        ['Being listed in "best agency" roundups', 'AI pulls from these listicles heavily (51% of AI citations)', 'TO DO'],
        ['FAQ content with direct Q&As', 'AI extracts direct answers from FAQ schemas', 'LIVE'],
        ['LinkedIn company page', 'Entity verification signal for all AI engines', 'TO DO'],
        ['Google Business Profile', 'Entity verification + local presence signal', 'TO DO'],
        ['3+ third-party reviews', 'Social proof signal — AI prioritises recommended vendors', 'TO DO'],
    ],
    col_widths=[2.4, 2.6, 0.8]
)

divider()

# ══════════════════════════════════════════════════════════════
# MEASUREMENT
# ══════════════════════════════════════════════════════════════
heading1('Measurement — What to Track')

add_table(
    ['Metric', 'Tool', 'Target by Day 30'],
    [
        ['Pages indexed in Google', 'Search Console', '15+ pages indexed'],
        ['Impressions (weekly)', 'Search Console', 'Trending up week-on-week'],
        ['Keyword positions', 'Ahrefs Webmaster Tools', '5+ keywords in top 50'],
        ['Backlinks acquired', 'Ahrefs Webmaster Tools', '10+ referring domains'],
        ['Directory citations live', 'Manual', '8+ directories'],
        ['Clutch reviews', 'Clutch profile', '3+ reviews'],
        ['AI mentions', 'Manual prompt testing', 'Appearing in 1+ AI engine'],
        ['Blog posts live', 'Website', '28+ posts (1/day from Day 1)'],
    ],
    col_widths=[2.0, 1.8, 2.3]
)

divider()

# ══════════════════════════════════════════════════════════════
# QUICK WINS
# ══════════════════════════════════════════════════════════════
heading1('Quick Wins — Do These First')
body('If time is limited, these five actions have the highest impact-to-effort ratio:')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.left_indent = Inches(0.25)
r1 = p.add_run('1.  Submit sitemap to Google Search Console and Bing Webmaster Tools')
set_font(r1, size=10.5, bold=True, color=BLACK)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(4)
p2.paragraph_format.left_indent = Inches(0.25)
r2 = p2.add_run('2.  Create Crunchbase and Clutch profiles today')
set_font(r2, size=10.5, bold=True, color=BLACK)

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(4)
p3.paragraph_format.left_indent = Inches(0.25)
r3 = p3.add_run('3.  Email 2-3 past clients asking for a Clutch review')
set_font(r3, size=10.5, bold=True, color=BLACK)

p4 = doc.add_paragraph()
p4.paragraph_format.space_after = Pt(4)
p4.paragraph_format.left_indent = Inches(0.25)
r4 = p4.add_run('4.  Search "best cold email agencies 2026" and email the top 5 article authors requesting inclusion')
set_font(r4, size=10.5, bold=True, color=BLACK)

p5 = doc.add_paragraph()
p5.paragraph_format.space_after = Pt(4)
p5.paragraph_format.left_indent = Inches(0.25)
r5 = p5.add_run('5.  Create a LinkedIn company page and post the signal-based outbound article today')
set_font(r5, size=10.5, bold=True, color=BLACK)

doc.add_paragraph()
body('Everything else in this plan builds on these five. The daily blog automation is already running — the content side is handled. The manual effort is the citation and link building work, and that is what moves rankings in weeks 2-4.')

divider()

p_final = doc.add_paragraph()
r_final = p_final.add_run('Lead Acquisition Ltd  ·  conor@leadacquisition.io  ·  leadacquisition.io')
set_font(r_final, size=9, color=LIGHT)

# Save
doc.save('C:/Users/cnrhs/Desktop/laio-website/Lead_Acquisition_30_Day_SEO_Plan.docx')
print('Done')
